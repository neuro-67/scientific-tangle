import fitz  # PyMuPDF
import docx  # python-docx
import json
import requests
import re
import argparse
import os
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ==========================================
# КОНФИГУРАЦИЯ API YANDEX CLOUD
# ==========================================
API_KEY = "НАСТЯ ВСТАВЬ КЛЮЧ СЮДА" 
FOLDER_ID = "НАСТЯ ВСТАВЬ АЙДИ ПАПКИ СЮДА"
TEXT_MODEL_URI = f"gpt://{FOLDER_ID}/yandexgpt/latest" 

# Сюда вставляем SYSTEM_PROMPT из текста выше
SYSTEM_PROMPT = """
Ты — Senior Data Architect в области горно-металлургического R&D.
Твоя задача: извлечь структурированный граф знаний из текста строго по утвержденной онтологии.
Выведи результат СТРОГО в формате JSON без пояснений.

ДОПУСТИМЫЕ ТИПЫ УЗЛОВ (label):
- Material (вещества, металлы, отходы)
- Process (технологии, методы)
- Equipment (оборудование)
- Property (параметр, например "температура", "сухой остаток")
- Measurement (конкретное числовое значение)
- Condition (условие, климат, режим)
- Experiment (опыт)
- Publication (отчет, статья)
- Expert (автор)
- Facility (завод, лаборатория)
- Finding (научный вывод, эффект)
- Topic (тег темы)

ДОПУСТИМЫЕ ТИПЫ СВЯЗЕЙ (type):
uses_material, applies_to, operates_at_condition, has_measurement, measures_property, uses_equipment, produces_output, showed, described_in, authored_by, expert_in, conducted_at, validated_by, contradicts, supports, tagged.

ПРАВИЛА ДЛЯ УЗЛОВ С ПАРАМЕТРАМИ (Measurement и Finding):
1. Если в тексте есть числа (концентрации, размеры, температура), создавай отдельный узел с label="Measurement". В его "properties" укажи: "value" (для точного числа), "min" (от/≥), "max" (до/≤), "unit" (размерность), "operator" (<=, >=, =, range).
2. Обязательно связывай Process/Experiment -> [has_measurement] -> Measurement -> [measures_property] -> Property.
3. ID для узлов Measurement и Finding генерируй как уникальные строки (например, "meas_sulf_200", "finding_yield_85").

Формат ответа:
{
  "nodes": [
    {"id": "обратный осмос", "label": "Process", "properties": {"domain": "экология"}},
    {"id": "шахтная вода", "label": "Material", "properties": {}},
    {"id": "meas_sulf_300", "label": "Measurement", "properties": {"max": 300, "unit": "мг/л", "operator": "<="}},
    {"id": "сульфаты", "label": "Property", "properties": {}}
  ],
  "edges": [
    {"source": "обратный осмос", "target": "шахтная вода", "type": "applies_to"},
    {"source": "обратный осмос", "target": "meas_sulf_300", "type": "has_measurement"},
    {"source": "meas_sulf_300", "target": "сульфаты", "type": "measures_property"}
  ]
}
"""

# ==========================================
# БЛОК ИЗВЛЕЧЕНИЯ ТЕКСТА (PDF + DOCX)
# ==========================================

def extract_text_from_pdf(file_path: str) -> str:
    print(f"[*] Открытие PDF файла: {file_path}")
    doc = fitz.open(file_path)
    full_text = ""
    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            full_text += f"\n\n--- СТРАНИЦА {page_num+1} ---\n\n" + text
    doc.close()
    return full_text

def extract_text_from_docx(file_path: str) -> str:
    print(f"[*] Открытие DOCX файла: {file_path}")
    doc = docx.Document(file_path)
    full_text = ""
    
    # Извлекаем обычные параграфы
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text.strip() + "\n"
            
    # Дополнительно извлекаем текст из таблиц (очень полезно для отчетов)
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_text += " | ".join(row_data) + "\n"
                
    return full_text

def extract_text(file_path: str) -> str:
    """Умный диспетчер, выбирающий парсер по расширению файла."""
    ext = file_path.lower().split('.')[-1]
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['doc', 'docx']:
        return extract_text_from_docx(file_path)
    else:
        print(f"[!] Формат .{ext} не поддерживается (пока что).")
        return ""

# ==========================================
# БЛОК LLM (NER)
# ==========================================

def extract_graph_from_text(text: str) -> dict:
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {
        "Authorization": f"Api-Key {API_KEY}", 
        "x-folder-id": FOLDER_ID, 
        "Content-Type": "application/json"
    }
    data = {
        "modelUri": TEXT_MODEL_URI,
        "completionOptions": {"stream": False, "temperature": 0.1, "maxTokens": "2000"},
        "messages": [
            {"role": "system", "text": SYSTEM_PROMPT},
            {"role": "user", "text": text}
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
             print(f"[!] Ошибка API YandexGPT: {response.status_code}")
             return {"nodes": [], "edges": []}
        
        result_text = response.json()['result']['alternatives'][0]['message']['text']
        
        match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        return {"nodes": [], "edges": []}
    except Exception as e:
        print(f"[!] Внутренняя ошибка парсинга: {e}")
        return {"nodes": [], "edges": []}

# ==========================================
# ГЛАВНЫЙ ПАЙПЛАЙН
# ==========================================
def process_file(file_path: str):
    if not os.path.exists(file_path):
        print(f"[!] Ошибка: Файл '{file_path}' не найден.")
        return

    file_name = os.path.basename(file_path)
    
    # 1. Извлекаем текст нужным парсером
    raw_text = extract_text(file_path)
    
    clean_text_length = len(raw_text.replace("\n", "").replace(" ", ""))
    MIN_CHARS_THRESHOLD = 200 
    
    if clean_text_length < MIN_CHARS_THRESHOLD:
        print(f"\n[!!!] ОШИБКА: Документ '{file_name}' не содержит читаемого текста (менее 200 символов). Пропущен.")
        return
        
    print(f"[*] Извлечено {clean_text_length} символов текста. Чанкинг...")
    
    # 2. Нарезка (TextSplitter)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunks = text_splitter.split_text(raw_text)
    print(f"[*] Получено {len(chunks)} частей.")
    
    all_nodes = []
    all_edges = []
    
    # 3. YandexGPT
    for i, chunk in enumerate(chunks):
        print(f"    - Обработка части {i+1} из {len(chunks)}...")
        graph_data = extract_graph_from_text(chunk)
        
        if "nodes" in graph_data:
            all_nodes.extend(graph_data["nodes"])
            
        if "edges" in graph_data:
            for edge in graph_data["edges"]:
                # РЕАЛИЗАЦИЯ ИЗ ТЗ (Модель провенанса): добавляем источник
                edge["properties"] = edge.get("properties", {})
                edge["properties"]["doc_id"] = file_name
            all_edges.extend(graph_data["edges"])
            
    final_graph = {
        "document": file_name,
        "nodes": all_nodes,
        "edges": all_edges
    }
    
    output_filename = f"{os.path.splitext(file_path)[0]}_graph.json"
    with open(output_filename, 'w', encoding='utf-8') as f:
        json.dump(final_graph, f, ensure_ascii=False, indent=2)
        
    print(f"\n[+] ГОТОВО! Сохранено: {output_filename}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Multi-format (PDF/DOCX) NER Pipeline для Графа Знаний.")
    parser.add_argument("filepath", type=str, help="Путь к PDF или DOCX файлу")
    args = parser.parse_args()
    process_file(args.filepath)