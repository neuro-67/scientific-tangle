import pymupdf # PyMuPDF
import docx  # python-docx
import json
import requests
import re
import argparse
import os
import base64
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ==========================================
# 1. КОНФИГУРАЦИЯ API YANDEX CLOUD
# ==========================================
API_KEY = "НАСТЯ ВСТАВЬ КЛЮЧ СЮДА" 
FOLDER_ID = "НАСТЯ ВСТАВЬ АЙДИ ПАПКИ СЮДА"

# Модель для сборки Графа (Онтология)
TEXT_MODEL_URI = f"gpt://{FOLDER_ID}/yandexgpt-5.1/latest" 
# Модель для чтения графиков (VLM)
VISION_MODEL_URI = f"gpt://{FOLDER_ID}/qwen3.6-35b-a3b/latest"

# ==========================================
# 2. СИСТЕМНЫЙ ПРОМПТ (СТРОГО ПО ONTOLOGY.md)
# ==========================================
SYSTEM_PROMPT = """
Ты — Senior Data Architect в области горно-металлургического R&D.
Твоя задача: извлечь структурированный граф знаний из текста строго по утвержденной онтологии.
Выведи результат СТРОГО в формате JSON без пояснений.

ДОПУСТИМЫЕ ТИПЫ УЗЛОВ (label):
- Material (вещества, металлы, отходы)
- Process (технологии, методы)
- Equipment (оборудование)
- Property (параметр, например "температура", "сухой остаток")
- Measurement (конкретное числовое значение)
- Condition (условие, климат, режим)
- Experiment (опыт)
- Publication (отчет, статья)
- Expert (автор)
- Facility (завод, лаборатория)
- Finding (научный вывод, эффект)
- Topic (тег темы)

ДОПУСТИМЫЕ ТИПЫ СВЯЗЕЙ (type):
uses_material, applies_to, operates_at_condition, has_measurement, measures_property, uses_equipment, produces_output, showed, described_in, authored_by, expert_in, conducted_at, validated_by, contradicts, supports, tagged.

ПРАВИЛА ДЛЯ УЗЛОВ С ПАРАМЕТРАМИ (Measurement и Finding):
1. Если в тексте есть числа (концентрации, размеры, температура), создавай отдельный узел с label="Measurement". В его "properties" укажи: "value" (для точного числа), "min" (от/≥), "max" (до/≤), "unit" (размерность), "operator" (<=, >=, =, range).
2. Обязательно связывай Process/Experiment -> [has_measurement] -> Measurement -> [measures_property] -> Property.
3. ID для узлов Measurement и Finding генерируй как уникальные строки (например, "meas_sulf_200", "finding_yield_85").

Формат ответа:
{
  "nodes": [
    {"id": "обратный осмос", "label": "Process", "properties": {"domain": "экология"}},
    {"id": "шахтная вода", "label": "Material", "properties": {}},
    {"id": "meas_sulf_300", "label": "Measurement", "properties": {"max": 300, "unit": "мг/л", "operator": "<="}},
    {"id": "сульфаты", "label": "Property", "properties": {}}
  ],
  "edges": [
    {"source": "обратный осмос", "target": "шахтная вода", "type": "applies_to"},
    {"source": "обратный осмос", "target": "meas_sulf_300", "type": "has_measurement"},
    {"source": "meas_sulf_300", "target": "сульфаты", "type": "measures_property"}
  ]
}
"""

# ==========================================
# 3. ФУНКЦИИ MULTIMODAL ПАРСИНГА
# ==========================================

def describe_page_visuals_with_qwen(image_bytes: bytes) -> str:
    """Отправляет картинку страницы в Qwen VL через OpenAI-совместимый API Яндекса."""
    encoded_image = base64.b64encode(image_bytes).decode('utf-8')
    
    # ВАЖНО: Для Qwen и Llama используется ЭНДПОИНТ OPENAI!
    url = "https://llm.api.cloud.yandex.net/v1/chat/completions"
    headers = {
        "Authorization": f"Api-Key {API_KEY}",
        # x-folder-id здесь не обязателен, папка зашита в URI модели
        "Content-Type": "application/json"
    }
    
    vision_prompt = (
        "Ты эксперт-металлург. Внимательно изучи эту страницу научно-технического отчета. "
        "Внимательно изучи эту страницу презентации. На ней есть результаты математического моделирования (CFD), "
        "цветные тепловые карты, схемы печей (например, Печь Ванюкова), графики скоростей, давлений или температур. "
        "Опиши максимально подробно всё, что видишь на картинках: какое оборудование показано, какие процессы, "
        "какие числовые значения на цветных шкалах (в мм, К, м/с). "
        "Извлеки все физические параметры. "
        "Если на странице НЕТ вообще никаких картинок, схем или цветовых карт, ответь 'ПУСТО'."
    )
    
    # Формат OpenAI для мультимодальных сообщений
    data = {
        "model": VISION_MODEL_URI, # Здесь ключ называется 'model', а не 'modelUri'
        "temperature": 0.1,
        "max_tokens": 1000,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": vision_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded_image}"}}
                ]
            }
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            resp_json = response.json()
            
            # Безопасно достаем content. Если его нет или он null, вернется None
            content = resp_json['choices'][0]['message'].get('content')
            
            # Если ответ пустой (None) или пустая строка
            if not content:
                return ""
                
            result = content.strip()
            
            if result.upper() == 'ПУСТО' or 'ПУСТО.' in result.upper():
                return ""
                
            return f"\n[ОПИСАНИЕ ГРАФИКОВ/СХЕМ СО СТРАНИЦЫ]: {result}\n"
        else:
            print(f"[!] Ошибка Qwen VL: {response.status_code} - {response.text}")
            return ""
    except Exception as e:
        print(f"[!] Ошибка вызова VLM: {e}")
        return ""


def extract_content_from_pdf(file_path: str) -> str:
    """Читает текст И отправляет скриншоты страниц в VLM."""
    print(f"[*] Открытие PDF: {file_path}")
    doc = pymupdf.open(file_path)
    full_text = ""
    
    for page_num, page in enumerate(doc):
        print(f"    - Страница {page_num+1} из {len(doc)}...")
        
        # 1. Достаем родной текстовый слой (работает мгновенно)
        native_text = page.get_text("text").strip()
        
        # 2. Делаем рендер страницы в картинку (150 dpi достаточно для 32B модели)
        pix = page.get_pixmap(dpi=150)
        image_bytes = pix.tobytes("jpeg")
        
        # 3. Отправляем в Qwen VL для анализа графиков
        print(f"      > Анализ графиков через Qwen2.5-VL...")
        visual_description = describe_page_visuals_with_qwen(image_bytes)
        
        if visual_description:
            print(f"      > Найдены и расшифрованы визуальные данные!")
            
        full_text += f"\n\n--- СТРАНИЦА {page_num+1} ---\n{native_text}\n{visual_description}"
        
    doc.close()
    return full_text

def extract_text_from_docx(file_path: str) -> str:
    """Чтение DOCX (без графиков, только текст и таблицы)."""
    print(f"[*] Открытие DOCX: {file_path}")
    doc = docx.Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        if para.text.strip(): full_text += para.text.strip() + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data: full_text += " | ".join(row_data) + "\n"
    return full_text

# ==========================================
# 4. БЛОК ГЕНЕРАЦИИ ГРАФА (YandexGPT)
# ==========================================

def extract_graph_from_text(text: str) -> dict:
    # НАСТОЯТЕЛЬНО рекомендую использовать PRO-версию для генерации сложного JSON
    PRO_MODEL_URI = f"gpt://{FOLDER_ID}/yandexgpt-5.1/latest" 
    
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {"Authorization": f"Api-Key {API_KEY}", "x-folder-id": FOLDER_ID, "Content-Type": "application/json"}
    data = {
        "modelUri": PRO_MODEL_URI,
        "completionOptions": {"stream": False, "temperature": 0.1, "maxTokens": "2000"},
        "messages": [{"role": "system", "text": SYSTEM_PROMPT}, {"role": "user", "text": text}]
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            result_text = response.json()['result']['alternatives'][0]['message']['text']
            match = re.search(r'\{.*\}', result_text, re.DOTALL)
            
            if match:
                json_string = match.group(0)
                try:
                    return json.loads(json_string)
                except json.JSONDecodeError as e:
                    print(f"\n[!!!] ОШИБКА JSON ОТ МОДЕЛИ: {e}")
                    print(f"Сырой ответ модели: {json_string[:300]}...\n")
                    return {"nodes": [], "edges": []}
            else:
                print(f"\n[!!!] МОДЕЛЬ НЕ ВЕРНУЛА JSON. Ответ: {result_text[:100]}...")
                return {"nodes": [], "edges": []}
        else:
            print(f"\n[!] Ошибка API YandexGPT: {response.text}")
    except Exception as e:
        print(f"\n[!] Критическая ошибка вызова LLM: {e}")
        
    return {"nodes": [], "edges": []}

# ==========================================
# 5. ГЛАВНЫЙ ПАЙПЛАЙН
# ==========================================
def process_file(file_path: str):
    if not os.path.exists(file_path):
        print(f"[!] Файл '{file_path}' не найден.")
        return

    file_name = os.path.basename(file_path)
    ext = file_path.lower().split('.')[-1]
    
    # 1. Извлечение контента (Текст + Графики)
    raw_text = extract_content_from_pdf(file_path) if ext == 'pdf' else extract_text_from_docx(file_path)
    
    if len(raw_text.replace("\n", "").replace(" ", "")) < 100:
        print(f"\n[!] ОШИБКА: Не удалось извлечь данные из '{file_name}'. Пропущен.")
        return
        
    print(f"\n[*] Контент извлечен. Начинаем разбивку (Chunking)...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=250)
    chunks = text_splitter.split_text(raw_text)
    
    all_nodes, all_edges = [], []
    
    # 2. Извлечение графа (NER)
    print(f"[*] Отправка данных в YandexGPT ({len(chunks)} частей)...")
    for i, chunk in enumerate(chunks):
        print(f"    - Генерация графа: часть {i+1} из {len(chunks)}...")
        graph_data = extract_graph_from_text(chunk)
        
        if "nodes" in graph_data:
            all_nodes.extend(graph_data["nodes"])
        if "edges" in graph_data:
            for edge in graph_data["edges"]:
                edge["properties"] = edge.get("properties", {})
                edge["properties"]["doc_id"] = file_name # Модель провенанса
            all_edges.extend(graph_data["edges"])
            
    # 3. Сохранение
    final_graph = {"document": file_name, "nodes": all_nodes, "edges": all_edges}
    out_name = f"{os.path.splitext(file_path)[0]}_graph.json"
    with open(out_name, 'w', encoding='utf-8') as f:
        json.dump(final_graph, f, ensure_ascii=False, indent=2)
        
    print(f"\n[+] УСПЕХ! Узлов: {len(all_nodes)}, Связей: {len(all_edges)}")
    print(f"[+] Файл: {out_name}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("filepath", type=str)
    args = parser.parse_args()
    process_file(args.filepath)